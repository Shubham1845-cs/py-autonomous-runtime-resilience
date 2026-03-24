"""
Generate AutoHeal-Py Workflow Documentation as DOCX
Run: python generate_workflow_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1, color=RGBColor(0x1a, 0x56, 0xdb)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1a56db")
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ═════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("AutoHeal-Py", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    run.font.size = Pt(28)

sub = doc.add_paragraph("Autonomous Self-Healing Microservices Framework")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()
sub2 = doc.add_paragraph("Complete Workflow Documentation")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub2.runs:
    run.font.size = Pt(12)
    run.font.bold = True

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  1. WHAT IS AUTOHEAL-PY
# ═════════════════════════════════════════════════════════════════════════════
heading("1. What is AutoHeal-Py?")
body(
    "AutoHeal-Py is a zero-touch, autonomous resilience framework for Python microservices. "
    "It monitors HTTP traffic in real time, detects failure patterns, and automatically injects "
    "the appropriate resilience pattern — all without modifying application code."
)
doc.add_paragraph()
body("Key Innovation:", bold=True)
bullet("No code changes required in the target service")
bullet("Autonomous detection and pattern injection via monkey-patching")
bullet("Three resilience patterns: Retry, Circuit Breaker, Timeout Guard")
bullet("Live dashboard with real-time topology visualization")
bullet("Self-healing: patterns are removed automatically when service recovers")

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  2. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
heading("2. System Architecture")
body(
    "AutoHeal-Py consists of four core components that form a closed-loop control system. "
    "Each component has a single responsibility and communicates through well-defined interfaces."
)
doc.add_paragraph()

add_table(
    headers=["Component", "File", "Responsibility"],
    rows=[
        ["TelemetryMonitor", "autoheal/monitor.py",
         "Monkey-patches requests library to intercept all HTTP calls and record metrics (latency, status codes, errors)"],
        ["HealthDetector",   "autoheal/detector.py",
         "Analyzes metrics using configurable thresholds to classify service health and recommend the correct resilience pattern"],
        ["PatternInjector",  "autoheal/injector.py",
         "Records active pattern injections, manages pattern lifecycle (inject / remove), and exposes state to the dashboard"],
        ["AutoHealAgent",    "autoheal/agent.py",
         "Orchestration loop: scans all services every 2s, calls Detector, calls Injector, emits events, and upgrades patterns when conditions escalate"],
    ],
    col_widths=[1.5, 1.8, 3.5]
)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  3. HOW IT WORKS — STEP BY STEP
# ═════════════════════════════════════════════════════════════════════════════
heading("3. How It Works — Step by Step")

steps = [
    ("Step 1: Install Monitor",
     "install_monitor() is called once at startup. It monkey-patches requests.get / post / put / delete / patch "
     "so every outgoing HTTP call is automatically intercepted and recorded in TelemetryMonitor's in-memory ring buffer."),
    ("Step 2: Traffic Generation",
     "A background thread sends real GraphQL queries to Saleor through the Fault Proxy (port 8001). "
     "The proxy can inject errors, delays, or connection drops on demand."),
    ("Step 3: Metric Collection",
     "For every HTTP call, TelemetryMonitor records: service name (from URL), response status code, "
     "latency in seconds, error message (if any), and timestamp."),
    ("Step 4: Health Analysis (every 2s)",
     "AutoHealAgent wakes up every 2 seconds and calls HealthDetector.analyze_health() for each service. "
     "The detector calculates failure rate and average latency over a 15-second sliding window."),
    ("Step 5: Pattern Selection",
     "HealthDetector._select_pattern() applies a priority-ordered decision tree:\n"
     "  Priority 1 — avg_latency > 3s → TIMEOUT GUARD\n"
     "  Priority 2 — failure_rate >= 60% → CIRCUIT BREAKER\n"
     "  Priority 3 — 503 error rate > 30% → RETRY"),
    ("Step 6: Pattern Injection",
     "If a pattern is recommended, AutoHealAgent calls PatternInjector.inject() to record the active pattern. "
     "The dashboard immediately reflects the new state. If conditions escalate (e.g., RETRY → CIRCUIT BREAKER), "
     "the agent automatically upgrades the pattern."),
    ("Step 7: Self-Healing",
     "When the fault is cleared, the service returns to healthy. After a 30-second grace period of sustained "
     "health, AutoHealAgent calls PatternInjector.remove() and the dashboard shows the service as HEALTHY again."),
]

for title_text, desc in steps:
    body(title_text, bold=True)
    body(desc)
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  4. THREE RESILIENCE PATTERNS
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Three Resilience Patterns")

add_table(
    headers=["Pattern", "Trigger Condition", "Dashboard Color", "What It Does"],
    rows=[
        ["Retry with Exponential Backoff",
         "503 error rate > 30% (DEGRADED state)",
         "Blue",
         "Retries failed requests up to 3 times with delays of 1s, 2s, 4s (2^n backoff) plus random jitter"],
        ["Circuit Breaker",
         "Failure rate >= 60% (CRITICAL state)",
         "Amber / Orange",
         "Opens the circuit after 5 failures — all subsequent calls fail immediately without hitting the service"],
        ["Timeout Guard",
         "Average latency > 3s (SLOW state)",
         "Purple",
         "Enforces a hard deadline on every call — threads that exceed the timeout are cancelled immediately"],
    ],
    col_widths=[1.8, 1.8, 1.2, 2.9]
)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  5. DEMO FLOW
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Live Demo Flow")
body(
    "The live_demo.py script walks through all three patterns in sequence. "
    "Each phase injects a specific fault via the Fault Proxy and waits for the agent to respond."
)
doc.add_paragraph()

add_table(
    headers=["Phase", "Fault Injected", "Expected Detection", "Expected Pattern"],
    rows=[
        ["Baseline",   "None",                    "HEALTHY",  "None"],
        ["Pattern 1",  "50% HTTP 503 errors",      "DEGRADED", "Retry with Exponential Backoff (Blue)"],
        ["Pattern 2",  "98% HTTP 503 errors",      "CRITICAL", "Circuit Breaker (Amber)"],
        ["Pattern 3",  "8-second response delay",  "SLOW",     "Timeout Guard (Purple)"],
        ["Recovery",   "Faults cleared",           "HEALTHY",  "None (pattern auto-removed after 30s)"],
    ],
    col_widths=[1.0, 1.8, 1.4, 2.5]
)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  6. DASHBOARD
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Dashboard Overview")
body("The Flask dashboard at http://localhost:5000 provides real-time visibility into the system.")
doc.add_paragraph()

body("Dashboard Pages:", bold=True)
bullet("/ (Home) — System overview: total services, total calls, active patterns, system status")
bullet("/dashboard — Live service cards with health state, failure rate, latency, active pattern badge")
bullet("/monitor — Detailed metrics table with per-call history")
bullet("/patterns — Pattern library: descriptions, configs, and use cases for all 3 patterns")
bullet("/docs — Full API reference and architecture documentation")

doc.add_paragraph()
body("Service Card States:", bold=True)
bullet("Green (HEALTHY) — No faults, no pattern active")
bullet("Yellow (DEGRADED) — Moderate failures, Retry pattern active (blue badge)")
bullet("Red (CRITICAL) — High failure rate, Circuit Breaker active (amber badge)")
bullet("Purple (SLOW) — High latency, Timeout Guard active (purple badge)")

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  7. HOW TO RUN
# ═════════════════════════════════════════════════════════════════════════════
heading("7. How to Run the Project")
body("You need 3 terminals running simultaneously:")
doc.add_paragraph()

add_table(
    headers=["Terminal", "Command", "Purpose"],
    rows=[
        ["Terminal 1", "cd AutoHeal-Py && python saleor_sandbox/fault_proxy.py",
         "Starts the fault injection proxy on port 8001"],
        ["Terminal 2", "cd AutoHeal-Py && python webapp/app.py",
         "Starts the Flask dashboard on port 5000 + AutoHeal agent"],
        ["Terminal 3", "cd AutoHeal-Py && python saleor_sandbox/live_demo.py",
         "Runs the interactive demo (press ENTER to advance each phase)"],
    ],
    col_widths=[1.0, 2.8, 2.9]
)

doc.add_paragraph()
body("Note:", bold=True)
body(
    "Saleor (port 8000) is optional — the system works without it. "
    "The fault proxy will return 502 errors for unresolvable upstream, but the monitor still "
    "tracks these as failures and the agent will still inject patterns correctly."
)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  8. DETECTION THRESHOLDS
# ═════════════════════════════════════════════════════════════════════════════
heading("8. Detection Thresholds (Configured in webapp/app.py)")

add_table(
    headers=["Threshold", "Value", "Triggers"],
    rows=[
        ["critical_failure_threshold", "60%",  "CRITICAL state → Circuit Breaker"],
        ["degraded_failure_threshold", "10%",  "DEGRADED state → Retry"],
        ["slow_latency_threshold",     "3.0s", "SLOW state → Timeout Guard"],
        ["scan_interval_seconds",      "2s",   "How often the agent scans all services"],
        ["grace_period_seconds",       "30s",  "How long service must be healthy before pattern removal"],
        ["Analysis window",            "15s",  "Sliding window for failure rate and latency calculation"],
    ],
    col_widths=[2.5, 1.0, 3.2]
)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
#  9. KEY INNOVATION
# ═════════════════════════════════════════════════════════════════════════════
heading("9. Key Innovation")
body(
    "AutoHeal-Py's core innovation is zero-touch resilience: the target application (Saleor) "
    "does not need to be modified in any way. The framework operates entirely at the HTTP client "
    "layer by monkey-patching Python's requests library."
)
doc.add_paragraph()
body("This means:", bold=True)
bullet("Any Python service using the requests library is automatically protected")
bullet("No SDK, no decorator, no configuration file needed in the target service")
bullet("The agent makes autonomous decisions — no human operator needed during runtime")
bullet("Patterns are dynamically upgraded as conditions worsen (Retry → Circuit Breaker)")
bullet("Full observability through the live dashboard without any instrumentation code")

doc.add_paragraph()
body(
    "The pattern selection algorithm uses a strict priority order to ensure the most appropriate "
    "pattern is always chosen: latency issues always trigger Timeout (not Circuit Breaker), "
    "because a slow service is fundamentally different from a failing service.",
    italic=True
)

# ═════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(os.path.dirname(__file__), "AutoHeal-Py_Workflow.docx")
doc.save(out_path)
print(f"✅ Saved: {out_path}")
